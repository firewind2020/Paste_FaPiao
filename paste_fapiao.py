import argparse
import glob
import os
import re
from datetime import datetime
from typing import Dict, List, Optional, Set, Tuple

import fitz  # PyMuPDF
import pandas as pd
import pdfplumber
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_BREAK
from docx.shared import Cm


DATE_RE = re.compile(r"开票日期[:：]\s*([0-9]{4}年[0-9]{2}月[0-9]{2}日)")
INVOICE_NO_RE = re.compile(r"发票号码[:：]\s*([0-9]{10,})")
BUYER_RE_RAIL = re.compile(r"购买方名称[:：]\s*(.*?)\s+统一社会信用代码")
BUYER_RE_COMMON = re.compile(r"购\s*名称[:：]\s*(.*?)\s+销\s*名称[:：]")
SELLER_RE_COMMON = re.compile(r"销\s*名称[:：]\s*(.*)$")

CITY_LIST_FILE = "中国大陆的民用定期航班通航城市.txt"
IATA_CODE_FILE = "中国国内航空公司的IATA航班代码.txt"
DEFAULT_STOP_WORDS = {
    "中国",
    "国内",
    "民用",
    "定期",
    "航班",
    "通航",
    "城市",
    "地区",
    "数据",
    "统计",
    "机场",
    "分布",
    "省份",
    "单位",
    "包含",
    "运营",
    "客运",
    "目前",
    "全国",
    "省份",
    "最多",
    "超",
    "个",
}

BUILTIN_CITY_DATA = """中国大陆的民用定期航班通航城市

根据中国民航局的数据（截至2024-2025年度统计），中国大陆的民用运输通航机场总数已达260多个。这些机场分布在以下城市和地区（以省份为单位划分，包含绝大多数正在运营客运航班的城市及地区）：

华北地区
北京：北京市（首都、大兴）
天津：天津市（滨海）
河北：石家庄、唐山、秦皇岛、邯郸、张家口、承德、邢台
山西：太原、大同、长治、运城、忻州、临汾、吕梁
内蒙古：呼和浩特、包头、海拉尔（呼伦贝尔）、赤峰、通辽、锡林浩特、乌海、鄂尔多斯、满洲里、乌兰浩特、巴彦淖尔、阿尔山、额济纳旗、扎兰屯、乌兰察布、霍林郭勒
东北地区
辽宁：沈阳、大连、丹东、锦州、朝阳、鞍山、营口
吉林：长春、延吉、长白山、通化、白城、松原
黑龙江：哈尔滨、齐齐哈尔、牡丹江、佳木斯、黑河、漠河、伊春、大庆、鸡西、加格达奇、抚远、建三江、五大连池、绥芬河
华东地区
上海：上海市（浦东、虹桥）
江苏：南京、无锡、徐州、常州、南通、连云港、淮安、盐城、扬州/泰州（共用）
浙江：杭州、宁波、温州、舟山、衢州、台州、义乌、丽水
安徽：合肥、黄山、阜阳、安庆、池州、芜湖/宣城（共用）
福建：福州、厦门、泉州、三明、武夷山、龙岩（连城）
江西：南昌、赣州、景德镇、九江、吉安、宜春、上饶、瑞金
山东：济南、青岛、烟台、威海、临沂、济宁、潍坊、东营、日照、菏泽
中南地区
河南：郑州、洛阳、南阳、信阳、安阳
湖北：武汉、宜昌、襄阳、恩施、神农架、十堰、荆州、鄂州
湖南：长沙、张家界、常德、怀化、永州、邵阳、岳阳、衡阳、郴州、湘西
广东：广州、深圳、珠海、揭阳（潮汕）、湛江、梅州、惠州、佛山、韶关
广西：南宁、桂林、北海、柳州、梧州、百色、河池、玉林
海南：海口、三亚、琼海、三沙
西南地区
重庆：重庆市（江北）、万州、黔江、武隆、巫山
四川：成都（双流、天府）、绵阳、泸州、宜宾、南充、达州、西昌、广元、攀枝花、巴中、稻城、康定、红原、九寨沟、甘孜
贵州：贵阳、遵义（新舟、茅台）、铜仁、毕节、兴义、安顺、荔波、黎平、黄平、六盘水
云南：昆明、丽江、西双版纳、芒市、大理、腾冲、迪庆、普洱、临沧、保山、昭通、文山、沧源、澜沧、宁蒗
西藏：拉萨、昌都、林芝、阿里、日喀则、隆子、定日、普兰
西北地区
陕西：西安、榆林、延安、汉中、安康
甘肃：兰州、敦煌、嘉峪关、庆阳、金昌、张掖、天水、陇南、甘南（夏河）
青海：西宁、格尔木、玉树、德令哈、花土沟、祁连、果洛
宁夏：银川、中卫、固原
新疆：（目前全国通航城市最多的省份，超20个）乌鲁木齐、喀什、和田、伊宁、库尔勒、阿克苏、吐鲁番、阿勒泰、库车、塔城、克拉玛依、哈密、博乐、布尔津、富蕴、石河子、莎车、若羌、图木舒克、于田、昭苏、阿拉尔、塔什库尔干、奇台、和静
"""

BUILTIN_IATA_DATA = """四大骨干航司
CA - 中国国际航空（Air China）
MU - 中国东方航空（China Eastern Airlines）
CZ - 中国南方航空（China Southern Airlines）
HU - 海南航空（Hainan Airlines）
地方主力与中型航司
3U - 四川航空（Sichuan Airlines）
ZH - 深圳航空（Shenzhen Airlines）
MF - 厦门航空（XiamenAir）
SC - 山东航空（Shandong Airlines）
FM - 上海航空（Shanghai Airlines）
JD - 首都航空（Capital Airlines）
GS - 天津航空（Tianjin Airlines）
HO - 吉祥航空（Juneyao Airlines）
低成本/民营及特色航司
9C - 春秋航空（Spring Airlines）
KN - 中国联合航空（China United Airlines）
G5 - 华夏航空（China Express Airlines）
GJ - 长龙航空（Loong Air）
8L - 祥鹏航空（Lucky Air）
PN - 西部航空（West Air）
EU - 成都航空（Chengdu Airlines）
TV - 西藏航空（Tibet Airlines）
BK - 奥凯航空（Okay Airways）
DZ - 东海航空（Donghai Airlines）
区域支线及其他航司
KY - 昆明航空（Kunming Airlines）
NS - 河北航空（Hebei Airlines）
QW - 青岛航空（Qingdao Airlines）
DR - 瑞丽航空（Ruili Airlines）
UQ - 乌鲁木齐航空（Urumqi Air）
GX - 北部湾航空（GX Airlines）
GY - 多彩贵州航空（Colorful Guizhou Airlines）
OQ - 重庆航空（Chongqing Airlines）
JR - 幸福航空（Joy Air）
CN - 大新华航空（Grand China Air）
Y8 - 金鹏航空（Suparna Airlines）
FU - 福州航空（Fuzhou Airlines）
9H - 长安航空（Air Changan）
GT - 桂林航空（Air Guilin）
A6 - 湖南航空（Air Travel，原红土航空）
RY - 江西航空（Jiangxi Air）
LT - 龙江航空（Longjiang Airlines）
9D - 天骄航空（Genghis Khan Airlines）
港澳台地区主要航司
CX - 国泰航空（Cathay Pacific）
HX - 香港航空（Hong Kong Airlines）
UO - 香港快运航空（HK Express）
HB - 大湾区航空（Greater Bay Airlines）
NX - 澳门航空（Air Macau）
CI - 中华航空（China Airlines）
BR - 长荣航空（EVA Air）
JX - 星宇航空（STARLUX Airlines）
"""


def read_pdf_text(pdf_path: str) -> str:
    parts: List[str] = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            parts.append(txt)
    return "\n".join(parts)


def classify_invoice(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    first_line = lines[0] if lines else ""
    has_rail_header = ("电子客票" in first_line) or ("电子客票" in text)
    has_rail_trip_pattern = bool(
        re.search(r"[GDCZTK]\d+", text)
        and re.search(r"[0-9]{4}年[0-9]{2}月[0-9]{2}日\s+[0-9]{2}:[0-9]{2}开", text)
    )
    if has_rail_header or has_rail_trip_pattern:
        return "高铁"

    detail_block = get_project_detail_block(lines)
    if "餐饮" in detail_block:
        return "餐饮"
    if "住宿" in detail_block:
        return "住宿"
    if "机票" in detail_block:
        return "飞机"
    if "民航发展基金" in detail_block:
        return "飞机"
    seller_name = extract_seller_name(text) or ""
    if "运输服务" in detail_block and "航空" in seller_name:
        return "飞机"

    # Fallback: try whole text if project-detail block is noisy.
    if "餐饮" in text:
        return "餐饮"
    if "住宿" in text:
        return "住宿"
    if "机票" in text:
        return "飞机"
    if "民航发展基金" in text:
        return "飞机"
    if "运输服务" in text and "航空" in seller_name:
        return "飞机"
    return "未知"


def is_special_invoice(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    first_line = lines[0] if lines else ""
    return "是" if "专用发票" in first_line else "否"


def get_project_detail_block(lines: List[str]) -> str:
    for idx, line in enumerate(lines):
        if "项目名称" in line:
            return " ".join(lines[idx + 1 : idx + 4])
    return ""


def extract_common_fields(text: str, invoice_type: str) -> Dict[str, Optional[str]]:
    result: Dict[str, Optional[str]] = {
        "开票日期": None,
        "发票号码": None,
        "购买方名称": None,
        "发票金额": None,
    }

    date_match = DATE_RE.search(text)
    if date_match:
        result["开票日期"] = date_match.group(1)

    no_match = INVOICE_NO_RE.search(text)
    if no_match:
        result["发票号码"] = no_match.group(1)

    buyer_match = BUYER_RE_RAIL.search(text) or BUYER_RE_COMMON.search(text)
    if buyer_match:
        result["购买方名称"] = buyer_match.group(1).strip()

    money_matches = re.findall(r"[¥￥]\s*([0-9]+\.[0-9]{2})", text)
    if money_matches:
        if invoice_type == "高铁":
            result["发票金额"] = money_matches[0]
        else:
            result["发票金额"] = money_matches[-1]

    return result


def extract_seller_name(text: str) -> Optional[str]:
    for line in text.splitlines():
        if "销" in line and "名称" in line:
            match = SELLER_RE_COMMON.search(line.strip())
            if match:
                return match.group(1).strip()
    return None


def extract_lodging_days(text: str, pdf_path: Optional[str] = None) -> Optional[str]:
    def _days_from_filename(name: str) -> Optional[int]:
        base = os.path.splitext(os.path.basename(name))[0]

        # Pattern A: 2025-12-05~06 (same year/month)
        match_a = re.search(
            r"([0-9]{4})[-_/年]([0-9]{1,2})[-_/月]([0-9]{1,2})\s*[~～至]\s*([0-9]{1,2})",
            base,
        )
        if match_a:
            year = int(match_a.group(1))
            month = int(match_a.group(2))
            day_start = int(match_a.group(3))
            day_end = int(match_a.group(4))
            try:
                dt_start = datetime(year, month, day_start)
                dt_end = datetime(year, month, day_end)
                if dt_end >= dt_start:
                    return (dt_end - dt_start).days + 1
            except ValueError:
                pass

        # Pattern B: 2025-12-05~2025-12-06
        match_b = re.search(
            r"([0-9]{4})[-_/年]([0-9]{1,2})[-_/月]([0-9]{1,2})\s*[~～至]"
            r"\s*([0-9]{4})[-_/年]([0-9]{1,2})[-_/月]([0-9]{1,2})",
            base,
        )
        if match_b:
            try:
                dt_start = datetime(int(match_b.group(1)), int(match_b.group(2)), int(match_b.group(3)))
                dt_end = datetime(int(match_b.group(4)), int(match_b.group(5)), int(match_b.group(6)))
                if dt_end >= dt_start:
                    return (dt_end - dt_start).days + 1
            except ValueError:
                pass

        return None

    if pdf_path:
        file_days = _days_from_filename(pdf_path)
        if file_days is not None:
            return str(file_days)

    for line in text.splitlines():
        if "住宿" in line:
            normalized_line = re.sub(r"\s+", " ", line).strip()

            # Quantity is usually the number before unit price/amount columns.
            qty_match = re.search(
                r"住宿(?:服务|费)?\*?\s+([0-9]+(?:\.[0-9]+)?)\s+[0-9]+(?:\.[0-9]+)?\s+[0-9]+\.[0-9]+",
                normalized_line,
            )
            if qty_match:
                try:
                    qty = float(qty_match.group(1))
                    # Guard against OCR-merging anomalies like 1328.98.
                    if 0 < qty <= 90:
                        return str(int(qty)) if qty.is_integer() else str(qty)
                except ValueError:
                    continue

    def _days_from_stay_dates(content: str) -> Optional[int]:
        default_year = datetime.now().year
        invoice_year_match = DATE_RE.search(content)
        if invoice_year_match:
            year_match = re.match(r"([0-9]{4})年", invoice_year_match.group(1))
            if year_match:
                default_year = int(year_match.group(1))

        def _extract_date_by_keywords(keywords: str) -> Optional[datetime]:
            full_patterns = [
                rf"(?:{keywords})[^0-9]{{0,12}}([0-9]{{4}})[年/-]([0-9]{{1,2}})[月/-]([0-9]{{1,2}})日?",
                rf"([0-9]{{4}})[年/-]([0-9]{{1,2}})[月/-]([0-9]{{1,2}})日?[^\n]{{0,12}}(?:{keywords})",
            ]
            for pattern in full_patterns:
                match = re.search(pattern, content)
                if match:
                    try:
                        return datetime(int(match.group(1)), int(match.group(2)), int(match.group(3)))
                    except ValueError:
                        continue

            short_patterns = [
                rf"(?:{keywords})[^0-9]{{0,12}}([0-9]{{1,2}})月([0-9]{{1,2}})日",
                rf"([0-9]{{1,2}})月([0-9]{{1,2}})日[^\n]{{0,12}}(?:{keywords})",
            ]
            for pattern in short_patterns:
                match = re.search(pattern, content)
                if match:
                    try:
                        return datetime(default_year, int(match.group(1)), int(match.group(2)))
                    except ValueError:
                        continue
            return None

        check_in = _extract_date_by_keywords(r"入住|入住日期|入住时间|住店|起住")
        check_out = _extract_date_by_keywords(r"离店|离店日期|退房|退住|离住")
        if check_in and check_out and check_out >= check_in:
            return max(1, (check_out - check_in).days)
        return None

    stay_days = _days_from_stay_dates(text)
    if stay_days is not None:
        return str(stay_days)

    return None


def normalize_city_name(name: str) -> str:
    cleaned = re.sub(r"（.*?）", "", name)
    cleaned = re.sub(r"\(.*?\)", "", cleaned)
    cleaned = cleaned.replace("市", "").strip()
    return cleaned


def load_city_set(base_dir: str) -> Set[str]:
    del base_dir

    city_set: Set[str] = set()
    for raw_line in BUILTIN_CITY_DATA.splitlines():
        line = raw_line.strip()
        if not line or "：" not in line:
            continue
        _, rhs = line.split("：", 1)
        for part in re.split(r"[、，,/\s]+", rhs):
            token = part.strip()
            if not token:
                continue
            token = re.sub(r"（.*?）", "", token)
            token = re.sub(r"\(.*?\)", "", token)
            if not re.fullmatch(r"[\u4e00-\u9fa5]{2,}", token):
                continue
            if token in DEFAULT_STOP_WORDS:
                continue

            normalized = normalize_city_name(token)
            if len(normalized) >= 2:
                city_set.add(normalized)
            if len(token) >= 2:
                city_set.add(token)
    return city_set


def load_iata_codes(base_dir: str) -> Set[str]:
    del base_dir

    codes: Set[str] = set()
    for raw_line in BUILTIN_IATA_DATA.splitlines():
        line = raw_line.strip()
        match = re.match(r"^([A-Z0-9]{2})\s*-", line)
        if match:
            codes.add(match.group(1))
    return codes


def extract_route_from_remark(lines: List[str], city_set: Set[str]) -> Tuple[Optional[str], Optional[str]]:
    for line in lines:
        for match in re.finditer(r"([\u4e00-\u9fa5]{2,})\s*[-—至]\s*([\u4e00-\u9fa5]{2,})", line):
            src = match.group(1).strip()
            dst = match.group(2).strip()
            src_n = normalize_city_name(src)
            dst_n = normalize_city_name(dst)

            src_ok = (src in city_set) or (src_n in city_set)
            dst_ok = (dst in city_set) or (dst_n in city_set)
            if src_ok and dst_ok:
                return src_n, dst_n
    return None, None


def extract_flight_no_from_remark(lines: List[str], iata_codes: Set[str]) -> Optional[str]:
    for line in lines:
        for match in re.finditer(r"\b([A-Z0-9]{2})(\d{3,4})\b", line):
            prefix = match.group(1)
            if not iata_codes or prefix in iata_codes:
                return f"{prefix}{match.group(2)}"
    return None


def extract_rail_fields(text: str) -> Dict[str, Optional[str]]:
    result: Dict[str, Optional[str]] = {
        "始发站": None,
        "终到站": None,
        "车次": None,
        "发车日期": None,
        "发车时间": None,
        "车厢座位号": None,
        "乘客姓名": None,
    }

    lines = [line.strip() for line in text.splitlines() if line.strip()]

    # Example: 北京南 G25 上海虹桥
    for line in lines:
        station_match = re.match(r"^(.*?)\s+([GDCZTK]\d+)\s+(.*?)$", line)
        if station_match:
            result["始发站"] = station_match.group(1).strip()
            result["车次"] = station_match.group(2).strip()
            result["终到站"] = station_match.group(3).strip()
            break

    # Example: 2026年03月08日 17:00开 07车05A号 二等座
    for line in lines:
        dt_match = re.search(r"([0-9]{4}年[0-9]{2}月[0-9]{2}日)\s+([0-9]{2}:[0-9]{2})开", line)
        if dt_match:
            result["发车日期"] = dt_match.group(1)
            result["发车时间"] = dt_match.group(2)

        seat_match = re.search(r"(\d+车\d+[A-Z]号)", line)
        if seat_match:
            result["车厢座位号"] = seat_match.group(1)

    # Example: 4403061969****0156 蓝洲
    for line in lines:
        name_match = re.search(r"\*{2,}\d*\s+([\u4e00-\u9fa5A-Za-z]+)$", line)
        if name_match:
            result["乘客姓名"] = name_match.group(1).strip()
            break

    return result


def extract_plane_fields(text: str, base_dir: str) -> Dict[str, Optional[str]]:
    result: Dict[str, Optional[str]] = {
        "航班号": None,
        "始发城市": None,
        "目的城市": None,
    }

    city_set = load_city_set(base_dir)
    iata_codes = load_iata_codes(base_dir)

    # Plane invoices usually place order/trip info in the remark area.
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    remark_lines: List[str] = []
    for line in lines:
        if "订单" in line or "备注" in line or "机票" in line or "运输服务" in line:
            remark_lines.append(line)
        elif "-" in line or "—" in line or "至" in line:
            remark_lines.append(line)

    if not remark_lines:
        remark_lines = lines

    result["航班号"] = extract_flight_no_from_remark(remark_lines, iata_codes)
    src, dst = extract_route_from_remark(remark_lines, city_set)
    result["始发城市"] = src
    result["目的城市"] = dst

    # Fallback when no external list is available.
    if not result["始发城市"] or not result["目的城市"]:
        for line in remark_lines:
            city_match = re.search(r"([\u4e00-\u9fa5]{2,})\s*[-—至]\s*([\u4e00-\u9fa5]{2,})", line)
            if city_match:
                result["始发城市"] = normalize_city_name(city_match.group(1))
                result["目的城市"] = normalize_city_name(city_match.group(2))
                break

    return result


def extract_plane_occurrence_date(text: str) -> Optional[str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    remark_lines: List[str] = []
    for line in lines:
        if "备注" in line or "订单" in line or "机票" in line or "运输服务" in line:
            remark_lines.append(line)
        elif "-" in line or "—" in line or "至" in line:
            remark_lines.append(line)

    if not remark_lines:
        remark_lines = lines

    # Pattern 1: 2025/12/11
    for line in remark_lines:
        date_match = re.search(r"([0-9]{4})/([0-9]{1,2})/([0-9]{1,2})", line)
        if date_match:
            year = int(date_match.group(1))
            month = int(date_match.group(2))
            day = int(date_match.group(3))
            return f"{year:04d}年{month:02d}月{day:02d}日"

    # Pattern 2: 12月13日 (use current year)
    for line in remark_lines:
        date_match = re.search(r"([0-9]{1,2})月([0-9]{1,2})日", line)
        if date_match:
            year = datetime.now().year
            month = int(date_match.group(1))
            day = int(date_match.group(2))

            # Year-crossing fix only for Pattern 2:
            # if invoice date is earlier than occurrence date, occurrence year should be previous year.
            try:
                occurrence_dt = datetime(year, month, day)
                invoice_date_match = DATE_RE.search(text)
                if invoice_date_match:
                    invoice_dt = datetime.strptime(invoice_date_match.group(1), "%Y年%m月%d日")
                    if invoice_dt < occurrence_dt:
                        year -= 1
            except ValueError:
                pass

            return f"{year:04d}年{month:02d}月{day:02d}日"

    return None


def parse_invoice(pdf_path: str) -> Dict[str, Optional[str]]:
    text = read_pdf_text(pdf_path)
    invoice_type = classify_invoice(text)

    row: Dict[str, Optional[str]] = {
        "文件名": os.path.basename(pdf_path),
        "发票类型": invoice_type,
        "是否专票": None,
        "开票日期": None,
        "发生日期": None,
        "发票号码": None,
        "购买方名称": None,
        "发票金额": None,
        "销售方名称": None,
        "住宿天数": None,
        "始发站": None,
        "终到站": None,
        "车次": None,
        "发车日期": None,
        "发车时间": None,
        "车厢座位号": None,
        "乘客姓名": None,
        "航班号": None,
        "始发城市": None,
        "目的城市": None,
    }

    row.update(extract_common_fields(text, invoice_type))
    row["是否专票"] = is_special_invoice(text)

    row["销售方名称"] = extract_seller_name(text)

    if invoice_type == "住宿":
        row["住宿天数"] = extract_lodging_days(text, pdf_path)
        row["发生日期"] = row.get("开票日期")

    if invoice_type == "高铁":
        row.update(extract_rail_fields(text))
        row["发生日期"] = row.get("发车日期")
        row["销售方名称"] = "中国国家铁路集团"

    if invoice_type == "餐饮":
        row["发生日期"] = row.get("开票日期")

    if invoice_type == "飞机":
        base_dir = os.path.dirname(pdf_path) or "."
        row.update(extract_plane_fields(text, base_dir))
        row["发生日期"] = extract_plane_occurrence_date(text) or row.get("开票日期")

    return row


def create_word_with_invoices(
    pdf_files: List[str],
    output_word: str,
    special_pdf_files: Optional[List[str]] = None,
) -> None:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0.4)
    section.right_margin = Cm(0.2)

    page_width_cm = section.page_width - section.left_margin - section.right_margin
    page_height_cm = section.page_height - section.top_margin - section.bottom_margin

    def _append_invoice_pages(invoice_files: List[str], temp_prefix: str) -> None:
        for i, _ in enumerate(invoice_files):
            if i % 2 == 0:
                if i > 0:
                    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
                    new_section.top_margin = Cm(0.5)
                    new_section.bottom_margin = Cm(0)
                    new_section.left_margin = Cm(0.4)
                    new_section.right_margin = Cm(0.2)

                table = doc.add_table(rows=2, cols=1)
                table.autofit = False
                table.width = page_width_cm

                for j in range(2):
                    idx = i + j
                    if idx < len(invoice_files):
                        cell = table.cell(j, 0)
                        cell.width = page_width_cm
                        cell.height = page_height_cm / 2
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                        try:
                            pdf_document = fitz.open(invoice_files[idx])
                            page = pdf_document.load_page(0)
                            pix = page.get_pixmap(dpi=300)
                            img_path = f"{temp_prefix}_{idx}.png"
                            pix.save(img_path)
                            pdf_document.close()

                            paragraph = cell.paragraphs[0]
                            paragraph.alignment = 1
                            run = paragraph.add_run()
                            run.add_picture(img_path, width=Cm(19.5))
                            os.remove(img_path)
                        except Exception as e:
                            print(f"处理 {invoice_files[idx]} 时出错: {e}")
                            cell.text = f"无法处理: {os.path.basename(invoice_files[idx])}"

    _append_invoice_pages(pdf_files, "temp_all")

    if special_pdf_files:
        page_break_para = doc.add_paragraph()
        page_break_para.add_run().add_break(WD_BREAK.PAGE)
        _append_invoice_pages(special_pdf_files, "temp_special")

    if os.path.exists(output_word):
        try:
            os.remove(output_word)
        except PermissionError:
            print(f"错误：请先关闭正在打开的 '{output_word}' 文件")
            raise

    doc.save(output_word)


def parse_cn_date_series(date_series: pd.Series) -> pd.Series:
    return pd.to_datetime(
        date_series.fillna("")
        .astype(str)
        .str.replace("年", "-", regex=False)
        .str.replace("月", "-", regex=False)
        .str.replace("日", "", regex=False),
        errors="coerce",
    )


def main() -> None:
    parser = argparse.ArgumentParser(description="提取发票信息并导出Excel")
    parser.add_argument(
        "--input-dir",
        default=".",
        help="PDF所在目录，默认当前目录",
    )
    parser.add_argument(
        "--output",
        default="发票信息提取结果.xlsx",
        help="输出Excel文件名",
    )
    parser.add_argument(
        "--word-output",
        default="发票粘贴.docx",
        help="输出Word文件名",
    )
    parser.add_argument(
        "--skip-word",
        action="store_true",
        help="仅导出Excel，不生成发票粘贴Word",
    )
    args = parser.parse_args()

    pdf_files = sorted(glob.glob(os.path.join(args.input_dir, "*.pdf")))
    if not pdf_files:
        raise FileNotFoundError("未找到PDF文件")

    rows = [parse_invoice(path) for path in pdf_files]
    df = pd.DataFrame(rows)

    if "发生日期" in df.columns:
        occurrence_dt = parse_cn_date_series(df["发生日期"])
        df = (
            df.assign(_发生日期排序=occurrence_dt)
            .sort_values(by="_发生日期排序", ascending=True, na_position="last", kind="mergesort")
            .drop(columns=["_发生日期排序"])
            .reset_index(drop=True)
        )

    output_path = os.path.join(args.input_dir, args.output)
    try:
        df.to_excel(output_path, index=False)
        print(f"已生成: {output_path}")
    except PermissionError:
        stem, ext = os.path.splitext(args.output)
        alt_name = f"{stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext or '.xlsx'}"
        alt_path = os.path.join(args.input_dir, alt_name)
        df.to_excel(alt_path, index=False)
        print(f"目标文件被占用，已生成: {alt_path}")

    if not args.skip_word:
        ordered_pdf_files = [os.path.join(args.input_dir, name) for name in df["文件名"].tolist()]
        special_ordered_pdf_files = [
            os.path.join(args.input_dir, name)
            for name in df.loc[df["是否专票"] == "是", "文件名"].tolist()
        ]
        word_output_path = os.path.join(args.input_dir, args.word_output)
        try:
            create_word_with_invoices(ordered_pdf_files, word_output_path, special_ordered_pdf_files)
            print(f"已生成: {word_output_path}")
        except PermissionError:
            stem, ext = os.path.splitext(args.word_output)
            alt_word_name = f"{stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext or '.docx'}"
            alt_word_path = os.path.join(args.input_dir, alt_word_name)
            create_word_with_invoices(ordered_pdf_files, alt_word_path, special_ordered_pdf_files)
            print(f"目标Word被占用，已生成: {alt_word_path}")


if __name__ == "__main__":
    main()
